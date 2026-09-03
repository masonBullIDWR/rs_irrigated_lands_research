'''
A script to calculate ET trends over time via OpenET Ensemble and NLCD data. This is an improvment upon the code entire_tv_et_trends.py
This will do the calculations and create a basic report as a word doc (hopefully)
'''
#%%set up cell, no calculations
import geopandas as gpd
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
import json
import ee, geemap
import re
from numpy import zeros
from matplotlib.ticker import FuncFormatter
from scipy import stats
from pathlib import Path

ee.Authenticate()
ee.Initialize(project= 'idwr-450722')
ee.data.setWorkloadTag('tv-et-trends')

first_year = 2005
final_year = 2024
#cdl data don't start in the TV until 2005
years_of_interest = list(range(first_year, final_year + 1))

et_version = 2.0
et_options = {2.0 : ['v2_0', '2.0'],
              2.1 : ['v2_1', '2.1']}

aoi_path = r"C:\Users\mason.bull\OneDrive - State of Idaho\Desktop\Geoprocessing\Data\TV\TV_and_MH_outline_union.shp"
aoi = gpd.read_file(aoi_path).to_crs('EPSG:8826')
aoi_ee = geemap.gdf_to_ee(aoi)

et = ee.ImageCollection(f"projects/openet/assets/ensemble/conus/gridmet/monthly/{et_options[et_version][0]}").filterBounds(aoi_ee).select('et_ensemble_mad')
eto = ee.ImageCollection(f"projects/openet/assets/reference_et/conus/gridmet/daily/v1").filterBounds(aoi_ee).select('eto')
cdl = ee.ImageCollection("USDA/NASS/CDL").filterBounds(aoi_ee).select('cropland')

#the mask for cdl is set up in 3 stages, minimal, medial, and agressive
#water is only masking out water, wetlands, and barren,
#forest is masking water, wetlands, barren, and forest 
#Desert is masking all of the above and desert areas
#all is masking everything that is not a crop
water = [0, 81, 83, 87, 88, 92, 111, 112, 190, 195]
forest = [0, 63, 81, 83, 87, 88, 92, 111, 112, 141, 142, 143, 190, 195]
desert = [0, 63, 64, 65, 81, 83, 87, 88, 92, 111, 112, 131, 141, 142, 143, 152, 176, 190, 195]
all = [0, 63, 64, 65, 81, 82, 83, 87, 88, 92, 111, 112, 121, 122, 123, 124, 131, 141, 142, 143, 152, 176, 190, 195]

#these masks filter out pixels across all images to return a masked image. ie., the water mask looks at all images 
#from 2000 to 2025 and masks out pixels that are water in ANY of the images, returning a single mask image
cdl_water_mask = cdl.map(lambda img: ee.Image(img).remap(water, [0]*len(water), 1)).sum().selfMask()

cdl_forest_mask = cdl.map(lambda img: ee.Image(img).remap(forest, [0]*len(forest), 1)).sum().selfMask()

cdl_desert_mask = cdl.map(lambda img: ee.Image(img).remap(desert, [0]*len(desert), 1)).sum().selfMask()

cdl_all_mask = cdl.map(lambda img: ee.Image(img).remap(all, [0]*len(all), 1)).sum().selfMask()

#a dictionary of values and colors for the cdl
cdl_lookup = json.load(open(r'C:\Users\mason.bull\OneDrive - State of Idaho\Desktop\Geoprocessing\Data\cdl_lookup.txt'))
colors = [cdl_lookup[c]['color'] for c in cdl_lookup]
colors_set = set(colors)
duplicates = [i for i in set(colors) if colors.count(i) > 1]
dashes = [(), (5,2), (1,3), (3,2,1,2), (4,1,2,1), (0.5,0.5)]
uses = {}
for i in cdl_lookup:
    item = cdl_lookup[i]
    color = item['color']
    if color in duplicates:
        try: 
            uses.update({color:uses[color]+1})
            item.update({'dash': dashes[uses[color]-1]})
        except:
            uses.update({color:0})
            uses.update({color:uses[color]+1})
            item.update({'dash': dashes[uses[color]-1]})
    else:
        item.update({'dash': ()})
#%%
#get data for ET and ETo via Earth Engine and export to Drive,
#this is a pretty large export, so the .start() of the export is commented out until you want to use it 
years_ee = ee.List(years_of_interest)

#function to unpack some dictionaries later
def unpackDict(di):
    keys = ee.Dictionary(di).keys()
    metric = keys.filter(ee.Filter.stringContains('item', 'crop').Not()).get(0)
    crop = ee.Dictionary(di).get('cropland')
    value = ee.Dictionary(di).get(metric)
    return ee.Dictionary([ee.String(crop), value])

#function to get growing season mean ET for each year over the aoi
def calculateET(year):
    year_str = ee.Number(year).format('%04d')
    filter_dates = ee.List([year_str.cat('-04-01'), year_str.cat('-11-01')])
    growing_season_cumm_et = et.filterDate(filter_dates.get(0), filter_dates.get(1)).sum()
    growing_season_cumm_eto = eto.filterDate(filter_dates.get(0), filter_dates.get(1)).sum()

    et_water_masked  = ee.Image(growing_season_cumm_et).updateMask(cdl_water_mask).rename(ee.String('et_water'))
    et_desert_masked = ee.Image(growing_season_cumm_et).updateMask(cdl_desert_mask).rename(ee.String('et_desert'))
    et_all_masked    = ee.Image(growing_season_cumm_et).updateMask(cdl_all_mask).rename(ee.String('et_all'))
    et_img = et_water_masked.addBands([et_desert_masked, et_all_masked])
    et_reductions = et_img.reduceRegion(reducer=ee.Reducer.mean(),
                                        geometry=aoi_ee.geometry(),
                                        scale=30,
                                        crs='EPSG:8826',
                                        maxPixels=3e7)
    
    eto_water_masked  = ee.Image(growing_season_cumm_eto).updateMask(cdl_water_mask).rename(ee.String('eto_water'))
    eto_desert_masked = ee.Image(growing_season_cumm_eto).updateMask(cdl_desert_mask).rename(ee.String('eto_desert'))
    eto_all_masked    = ee.Image(growing_season_cumm_eto).updateMask(cdl_all_mask).rename(ee.String('eto_all'))
    eto_img = eto_water_masked.addBands([eto_desert_masked, eto_all_masked])
    eto_reductions = eto_img.reduceRegion(reducer=ee.Reducer.mean(),
                                        geometry=aoi_ee.geometry(),
                                        scale=30,
                                        crs='EPSG:8826',
                                        maxPixels=3e7)

    #now get landcover area per crop and crop et
    cdl_img = ee.ImageCollection(cdl.filterDate(year_str)).first()

    et_lc_img = ee.Image(growing_season_cumm_et).addBands(cdl_img)

    lc_et_reducer = ee.Reducer.mean().group(groupField= 1, groupName='cropland')
    lc_area_reducer = ee.Reducer.count().group(groupField= 1, groupName='cropland')
    lc_et = ee.Image(et_lc_img).reduceRegion(reducer= lc_et_reducer.combine(
                                                                      lc_area_reducer, 
                                                                      outputPrefix='area', 
                                                                      sharedInputs=True),
                                       geometry=aoi_ee.geometry(),
                                       scale=30,
                                       crs='EPSG:8826',
                                       maxPixels=3e7)
    crop_area = lc_et.get('areagroups')
    crop_et = lc_et.get('groups')
    crop_et_dict = ee.List(crop_et).map(unpackDict)
    crop_area_dict = ee.List(crop_area).map(unpackDict)
        
    return ee.Feature(None, {'year': year, 'crop_area': crop_area_dict, 'crop_et': crop_et_dict,
                             'et_water': et_reductions.get('et_water'), 'et_desert': et_reductions.get('et_desert'), 'et_all': et_reductions.get('et_all'), 
                             'eto_water': eto_reductions.get('eto_water'), 'eto_desert': eto_reductions.get('eto_desert'), 'eto_all': eto_reductions.get('eto_all'),
                             })

entire_aoi_et_values = ee.FeatureCollection(years_ee.map(calculateET))

#there has to be an export here because the reductions are so large
ee.batch.Export.table.toDrive(collection=entire_aoi_et_values,
                              description='tv_et_lc_trends_export',
                              fileNamePrefix='tv_et_lc_trends',
                              )#.start() #this gets commented out as a failsafe to not accidentally export the csv 

#%%
#getting the csv into the format we want before plotting
#this is the csv that is exported from EE
et_data = pd.read_csv(r"C:\Users\mason.bull\OneDrive - State of Idaho\Desktop\Geoprocessing\Data\TV\et_trends\tv_et_lc_trends.csv")
et_data = et_data.drop(columns=['.geo', 'system:index'])

#this gets us a row id to separate lines on instead of a year, incase there is an empty year
et_data = et_data.reset_index(drop = True).rename_axis('row_id').reset_index()

#the expression that is in each string that gets output by the EE export
et_expression = re.compile(r'(\w+)=([\d$.]+)')

#function to find area or et values by crop, organize them, and neatly store them in a dataframe
def parse_metrics(metric):
    column_name = f'crop_{metric}'

    items = [(row_id, int(k), float(v))
             for row_id, s in zip(et_data['row_id'], et_data[column_name])
             for k, v in et_expression.findall(s)]
    return pd.DataFrame(items, columns = ['row_id', 'crop_code', column_name])

et_long = parse_metrics('et')
area_long = parse_metrics('area')

#combine the crop type dfs and add a column for crop name and color code
crop_long = et_long.merge(area_long, on = ['row_id', 'crop_code'], how = 'outer')
crop_long['crop'] = crop_long['crop_code'].map(lambda c: cdl_lookup.get(str(c), {}).get('crop', f'unknown_{c}'))
crop_long['color_code'] = crop_long['crop_code'].map(lambda c: cdl_lookup.get(str(c), {}).get('color', f'unknown_{c}'))
crop_long['dash_code'] = crop_long['crop_code'].map(lambda c: cdl_lookup.get(str(c), {}).get('dash', f'unknown_{c}'))

columns_to_pivot = et_data[['row_id', 'year', 'et_all', 'et_water', 'et_desert', 'eto_all', 'eto_water', 'eto_desert']]

#use a single pivot to get the data to a long format without having to stack up melt arguments
filters_long = pd.wide_to_long(columns_to_pivot, stubnames=['et', 'eto'], i='row_id', 
                               j = 'filter', sep='_', suffix=(r'\w+')).reset_index()

#finalize and add new columns we want
df_final = crop_long.merge(filters_long, on='row_id', how='inner')
df_final['et_of'] = df_final['et']/df_final['eto']
df_final['year'] = df_final['year'].astype(int)
df_final['crop_area'] = df_final['crop_area']*9e-4

for i in ['Water', 'Clouds/No Data', 'Grass/Pasture', 'Shrubland', 'Open Water', 'Perennial Ice/Snow',
          'Forest', 'Developed', 'Deciduous Forest', 'Evergreen Forest', 'Woody Wetlands', 'Herbaceous Wetlands'
          'Mixed Forest', 'Wetlands']: #pick lc types to exclude from plotting
    df_final = df_final[df_final.crop != i]

#make a palette for plotting crop types
palette = {}
for k in cdl_lookup:
    palette.update({cdl_lookup[k]['crop']:cdl_lookup[k]['color']})

dash_palette = {}
for k in cdl_lookup:
    dash_palette.update({cdl_lookup[k]['crop']:cdl_lookup[k]['dash']})

#find the ten largest lc types by area for each year to help clean up plots
top_ten_set = set()
for y in years_of_interest:
    df = df_final[df_final.year == y]
    df = df[df['filter'] == 'all']
    largest = pd.DataFrame(df['crop_area'].nlargest(10))
    for i in set(df_final.loc[list(largest.index)]['crop']):
        top_ten_set.add(i)

df_final_top_ten = df_final.loc[df_final.crop.isin(top_ten_set)]

def linearRegression(ind, dep):
    slope, intercept, r_value, p_value, std_err = stats.linregress(df_final[df_final['filter'] == 'all'][ind], df_final[df_final['filter'] == 'all'][dep])
    return slope, intercept, r_value, p_value, std_err
def getPReport(p_value):
    if p_value < 0.001:
        p = 'p < 0.001'
    elif p_value < 0.01 and p_value > 0.001:
        p = 'p < 0.01'
    elif p_value < 0.05 and p_value >0.01:
        p = 'p < 0.05'
    else:
        p = f'p = {p_value:.3f}'
    return p
#%%
#plotting

#plot of ET depth 
fig, ax = plt.subplots()
slope, intercept, r_value, p_value, std_err = linearRegression('year', 'et')
#regplot give the trendline of the data, it only seems necessary to include it for one type of data
sns.regplot(data=df_final[df_final['filter'] == 'all'], x = 'year', y = 'et', color='blue', scatter = False, line_kws={'linestyle': 'dashed'})
#sns.regplot(data=df_final[df_final['filter'] == 'desert'], x = 'year', y = 'et', color='green', scatter = False)
#sns.regplot(data=df_final[df_final['filter'] == 'water'], x = 'year', y = 'et', color='orange', scatter = False)
sns.lineplot(data=df_final, x = 'year', y = 'et', hue='filter')
ax.text(2007, 540, f'r² = {r_value**2:.2f}\nslope = {slope:.2f}, {getPReport(p_value)}') #scipy reports the r value, need to square it for reporting
ax.set_title('Treasure Valley ET')
ax.set_xlim(first_year, final_year)
ax.set_ylabel('ET (mm)')
ax.set_xlabel('Year')
handles, labels = ax.get_legend_handles_labels()
ax.legend(title = None, handles = handles, labels = ['Only Crops', 'No Water', 'No Water or Desert'])
sns.move_legend(ax, 'upper center', bbox_to_anchor = (0.5,-0.1), ncols = 3)
plt.gca().xaxis.set_major_formatter(FuncFormatter(lambda x,_: int(x)))
fig.savefig('et_depth.png', bbox_inches = 'tight')
fig.show()

#%%
def makePlot(dep, plot_focus, y_axis_lab, ET_plot, plot_file_name, text_x = None, text_y = None, dataframe = df_final, ind= 'year', show = False, save = False, y_lim = (None, None)):
    fig, ax = plt.subplots()
    if ET_plot:
        sns.lineplot(data=dataframe, x = ind, y = dep, hue='filter')
        slope, intercept, r_value, p_value, std_err = linearRegression(ind, dep)
        #regplot give the trendline of the data, it only seems necessary to include it for one type of data
        sns.regplot(data=dataframe[dataframe['filter'] == 'all'], x = ind, y = dep, color='blue', scatter = False, line_kws={'linestyle': 'dashed'})
        #sns.regplot(data=df_final[df_final['filter'] == 'desert'], x = 'year', y = 'et', color='green', scatter = False)
        #sns.regplot(data=df_final[df_final['filter'] == 'water'], x = 'year', y = 'et', color='orange', scatter = False)
        ax.text(text_x, text_y, f'r² = {r_value**2:.2f}\nslope = {slope:.2f}, {getPReport(p_value)}') #scipy reports the r value, need to square it for reporting
        ax.set_title(f'Treasure Valley {plot_focus}')
    else:
        sns.lineplot(data = dataframe, x = 'year', y = 'crop_area', palette= palette, hue = 'crop', style = 'crop', dashes = dash_palette)
        ax.set_title(plot_focus)
    ax.set_xlim(first_year, final_year)
    ax.set_ylim(y_lim)
    ax.set_ylabel(y_axis_lab)
    ax.set_xlabel('Year')
    handles, _ = ax.get_legend_handles_labels()
    if ET_plot:
        ax.legend(title = None, handles = handles, labels = ['Only Crops', 'No Water', 'No Water or Desert'])
    sns.move_legend(ax, 'upper center', bbox_to_anchor = (0.5,-0.1), ncols = 3)
    plt.gca().xaxis.set_major_formatter(FuncFormatter(lambda x,_: int(x)))
    if save:
        fig.savefig(f'{plot_file_name}.png', bbox_inches = 'tight')
        print(f'Plot saved to {str(Path.cwd()/plot_file_name)}.png')
    if show:
        fig.show()

makePlot(dep = 'et', plot_focus='ET', y_axis_lab='ET (mm)', ET_plot=True, 
         plot_file_name='et_depth', text_x = 2005.5, text_y=550, save=True)
makePlot('eto', 'ETo', 'ETo (mm)', True, 'eto', 2005.5, 1080, save= True)
makePlot('et_of', 'EToF', 'EToF ', True, 'etof', 2005.5, 0.556, save= True)
makePlot('crop_area', 'Crop Area', 'Area (km²)', False, 'crop_area', save= True)

makePlot('crop_area', 'Crop Area (top ten crops annually)', 'Area (km²)', 
         False, 'top_ten_crop_area', dataframe=df_final_top_ten, y_lim=(0, 900), save= True)
makePlot('crop_area', 'Crop Area (zoomed, top ten crops annually)', 'Area (km²)', 
         False, 'top_ten_crop_area_zoomed', dataframe=df_final_top_ten, y_lim=(0, 100), save= True)
makePlot('crop_et', 'Crop ET (top ten crops annually)', 'ET (mm)', 
         False, 'top_ten_crop_et', dataframe=df_final_top_ten, save= True)

#%%
#plot of ETo
fig, ax = plt.subplots()
sns.lineplot(data=df_final, x = 'year', y = 'eto', hue='filter')
ax.set_title('Treasure Valley ETo')
ax.set_xlim(first_year, final_year)
ax.set_ylabel('ETo (mm)')
ax.set_xlabel('Year')
handles, labels = ax.get_legend_handles_labels()
ax.legend(title = None, loc = 'center right', handles = handles, labels = ['Only Crops', 'No Water', 'No Water or Desert'])
ax.legend(title = None, loc = 'lower right')
plt.gca().xaxis.set_major_formatter(FuncFormatter(lambda x,_: int(x)))
fig.savefig('eto.png', bbox_inches = 'tight')
fig.show()

#plot of EToF
fig, ax = plt.subplots()
sns.lineplot(data=df_final, x = 'year', y = 'et_of', hue='filter')
ax.set_title('Treasure Valley EToF')
ax.set_xlim(first_year, final_year)
ax.set_ylabel('EToF')
ax.set_xlabel('Year')
handles, labels = ax.get_legend_handles_labels()
ax.legend(title = None, loc = 'center right', handles = handles, labels = ['Only Crops', 'No Water', 'No Water or Desert'])
ax.legend(title = None, loc = 'lower right')
plt.gca().xaxis.set_major_formatter(FuncFormatter(lambda x,_: int(x)))
fig.savefig('etof.png', bbox_inches = 'tight')
fig.show()

#plot of area for each crop 
fig, ax = plt.subplots()
sns.lineplot(data = df_final, x = 'year', y = 'crop_area', palette= palette, hue = 'crop', style = 'crop', dashes = dash_palette)
ax.set_xlim(first_year, final_year)
ax.set_title('Area per crop')
ax.set_ylabel('Area (km²)')
ax.set_xlabel('Year')
ax.legend(title = None)
sns.move_legend(ax, 'upper center', bbox_to_anchor = (0.5,-0.1), ncols = 3)
plt.gca().xaxis.set_major_formatter(FuncFormatter(lambda x,_: int(x)))
fig.savefig('crop_area.png', bbox_inches = 'tight')
fig.show()

#plot of ET for each crop 
fig, ax = plt.subplots()
sns.lineplot(data = df_final, x = 'year', y = 'crop_et', palette= palette, hue = 'crop', style = 'crop', dashes = dash_palette)
ax.set_xlim(first_year, final_year)
ax.set_title('ET depth per crop')
ax.set_ylabel('Depth (mm)')
ax.set_xlabel('Year')
ax.legend(title = None)
sns.move_legend(ax, 'upper center', bbox_to_anchor = (0.5,-0.1), ncols = 3)
plt.gca().xaxis.set_major_formatter(FuncFormatter(lambda x,_: int(x)))
fig.savefig('crop_et.png', bbox_inches = 'tight')
fig.show()

#these are plots of the largest crops by area, top ten in each year 
df_final_top_ten = df_final.loc[df_final.crop.isin(top_ten_set)]

#plot of area for each crop 
fig, ax = plt.subplots()
sns.lineplot(data = df_final_top_ten, x = 'year', y = 'crop_area', palette= palette, hue = 'crop', style = 'crop', dashes = dash_palette)
ax.set_xlim(first_year, final_year)
ax.set_ylim(0, 900)
ax.set_title('Area per crop (top ten crops by area annually)')
ax.set_ylabel('Area (km²)')
ax.set_xlabel('Year')
sns.move_legend(ax, 'upper center', bbox_to_anchor = (0.5,-0.1), ncols = 3)
plt.gca().xaxis.set_major_formatter(FuncFormatter(lambda x,_: int(x)))
fig.savefig('top_ten_crop_area.png', bbox_inches = 'tight')
fig.show()

fig, ax = plt.subplots()
sns.lineplot(data = df_final_top_ten, x = 'year', y = 'crop_area', palette= palette, hue = 'crop', style = 'crop', dashes = dash_palette)
ax.set_xlim(first_year, final_year)
ax.set_ylim(0, 100)
ax.set_title('Area per crop (zoomed, top ten crops by area annually)')
ax.set_ylabel('Area (km²)')
ax.set_xlabel('Year')
sns.move_legend(ax, 'upper center', bbox_to_anchor = (0.5,-0.1), ncols = 3)
plt.gca().xaxis.set_major_formatter(FuncFormatter(lambda x,_: int(x)))
fig.savefig('top_ten_crop_area_zoomed.png', bbox_inches = 'tight')
fig.show()

#plot of ET for each crop 
fig, ax = plt.subplots()
sns.lineplot(data = df_final_top_ten, x = 'year', y = 'crop_et', palette= palette, hue = 'crop', style = 'crop', dashes = dash_palette)
ax.set_xlim(first_year, final_year)
ax.set_title('ET depth per crop (top ten crops by area annually)')
ax.set_ylabel('Depth (mm)')
ax.set_xlabel('Year')
ax.legend(title = None)
sns.move_legend(ax, 'upper center', bbox_to_anchor = (0.5,-0.1), ncols = 3)
plt.gca().xaxis.set_major_formatter(FuncFormatter(lambda x,_: int(x)))
fig.savefig('top_ten_crop_et.png', bbox_inches = 'tight')
fig.show()

#%%Create the document with figures and text
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from datetime import datetime
date = datetime.now().strftime('%Y-%m-%d')
filters_dict = {'Water': [water, "'Water' is the least aggressive filter, "],
                'Desert': [desert, "'Desert' is the more aggressive filter, "],
                'All': [all, "'All' is the most aggressive filter, "]}

for n in filters_dict:
    item_names = [cdl_lookup[str(i)]['crop'] for i in filters_dict[n][0]]
    item_names_end = item_names[-1]
    item_names.remove(item_names_end)
    item_names.insert(len(item_names) + 1, f'and {item_names_end}')
    item_string = f'{filters_dict[n][1]}filtering out {', '.join(item_names)}.\n'
    filters_dict[n].append(item_string)

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
doc.add_heading('Treasure Valley ET Trends', 0)
doc.add_heading(f'Created by Mason Bull on {date}', 2)

doc.add_heading('Background', 1)
doc.add_paragraph(f'Data were created in GEE and ripped down for plotting. We used ET measurments from OpenET Ensemble Version {et_version}.' \
                  f'We gathered landcover information from the NASS CDL data. Due to data availability we can map ET trends as far back as {first_year}.')

doc.add_heading('Methods', 1)
p = doc.add_paragraph(f"ET data were first summed per pixel for each month between April and October for each year of the analysis." \
                  "The summed growing season ET was then averaged across the study area. The study area is the boundary of the TVGWFM, from Glenns Ferry to Payette, with the western boundary of the Snake River, and East boundary of Long Tom and Black Canyon reservoirs." \
                  "We also gathered ETo data from OpenET to calculate EToF across the region. ET data were filtered by three groups of landcover before the averaging calculation.\n")
doc.add_paragraph(f"The first filtering group is referred to as 'Water'. {filters_dict['Water'][-1]}", style='List Bullet') 
doc.add_paragraph(f"The first filtering group is referred to as 'Desert'. {filters_dict['Desert'][-1]}", style='List Bullet')
doc.add_paragraph(f"The third filtering group is referred to as 'All'. {filters_dict['All'][-1]}", style='List Bullet')
p.add_run(f"We then calculated ET and area of each crop type in the study area per CDL classification.")

doc.add_heading('Figures', 1)
doc.add_picture('et_depth.png')
doc.add_paragraph(f"Figure 1. ET rate (mm) across the Treasure Valley from {first_year} to {final_year}. Blue represents the heavily filtered 'All' class. Green represents the moderately filtered 'Desert' class. Orange represents the lightly filtered 'Water' class. Dashed blue line represents the linear regression of et of crops in the Treasure Valley.")

doc.add_picture('eto.png')
doc.add_paragraph(f"Figure 2. ETo depth (mm) across the Treasure Valley from {first_year} to {final_year}. Blue represents the heavily filtered 'All' class. Green represents the moderately filtered 'Desert' class. Orange represents the lightly filtered 'Water' class.")

doc.add_picture('etof.png')
doc.add_paragraph(f"Figure 3. EToF (unitless) across the Treasure Valley from {first_year} to {final_year}. Blue represents the heavily filtered 'All' class. Green represents the moderately filtered 'Desert' class. Orange represents the lightly filtered 'Water' class.")

doc.add_picture('top_ten_crop_area.png')
doc.add_paragraph(f"Figure 4. Landcover area of the ten largest landcover classes in km² across the Treasure Valley from {first_year} to {final_year}. Shrubland, water, and grass/pasture are excluded. Colors are the CDL landcover class color codes.")

doc.add_picture('top_ten_crop_area_zoomed.png')
doc.add_paragraph(f"Figure 5. Landcover area of the ten largest landcover classes in km² across the Treasure Valley from {first_year} to {final_year}. Zoomed in to show detail of crops with <= 100 km² Shrubland, water, and grass/pasture are excluded. Colors are the CDL landcover class color codes.")

doc.add_picture('top_ten_crop_et.png')
doc.add_paragraph(f"Figure 6. ET depth (mm) of the ten largest landcover classes in km² across the Treasure Valley from {first_year} to {final_year}. Shrubland, water, and grass/pasture are excluded. Colors are the CDL landcover class color codes.")

doc.add_heading('Discussion', 1)

doc.save('report.docx')
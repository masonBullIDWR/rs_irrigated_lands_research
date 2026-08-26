#%%
from pathlib import Path
from shutil import copy
from arcpy import metadata, mp, SpatialReference
from PIL import Image, ImageDraw, ImageFont
from os.path import getmtime
from time import strftime, strptime, ctime, time
import xml.etree.ElementTree as ET
from docx import Document
from python_docx_replace import docx_replace
from ruamel.yaml import YAML
import json
from thumbnail_generation import generateThumbnail


#-----------------------static variables -------------------------
#template metadata file for classified imagery
template_xml = r"X:\Spatial\LandCover_Vegetation\SnakePlain\MachineLearning\ESPA_2024_RandomForest.tif.xml"

#get the configuration info
parent_dir = Path.cwd().parent.absolute()
config_file = [i for i in parent_dir.glob('*.yml')][0]
yaml = YAML()
yaml.preserve_quotes = (True)
with open(config_file) as f:
    config = yaml.load(f)

json_path = [i for i in parent_dir.glob('**/metadata_dictionaries.json')][0]
with open(json_path) as js:
    file = json.load(js)
    location_dict = file['location_dict']
    datasets_dict = file['datasets_dict']
    long_metadata = file['long_metadata']

year = str(config['year'])

#region is a key lookup value for the location_dict (the N: location)
area = config['area']

#this if is a weird edge case with the renaming of the Treasure Valley and how we are handling that moving forward
if area != 'tv':
    region = [l for l in location_dict.keys() if area.casefold() == str(location_dict[l][1]).casefold()][0]
else:
    region = 'TreasureValley'

#path to root folder of training data to find the reporting folder
root_path = Path(config['training_data']).parent.parent

x_drive_name = location_dict[region][0] #ie., BoiseValley
abb_name = location_dict[region][1] #ie., TV
full_name = location_dict[region][2] #ie., Treasure Valley

#----------------file locations--------------------
n_loc = f'N:\\IrrigatedLands\\{region}\\RandomForest_{year}\\forRelease'
x_spatial_loc = f'X:\\Spatial\\LandCover_Vegetation\\{x_drive_name}\\MachineLearning'
x_staging_loc = f'X:\\Staging_X_Y\\LandCover_Vegetation\\{x_drive_name}\\MachineLearning'
metadata_loc = f'N:\\IrrigatedLands\\rf_metadata_template.docx'

#metadata document in a docx format for easy editing when things need changed
metadata_doc = Document(metadata_loc)

if not Path(Path.cwd()/'temp').exists():
    Path(Path.cwd()/'temp').mkdir()
temp_folder = str(Path.cwd()/'temp')

#make the thumbnail for the portal item 
generateThumbnail(year, full_name, temp_folder, n_loc)

#----------------file setup------------------------
def setupDirectories(to_location = x_staging_loc, from_location = Path(n_loc)):
    '''Get the directories on public folders set up, rename items, and copy data 
    to where it needs to go.'''
    if not Path(to_location).exists():
        Path(to_location).mkdir(parents=True, exist_ok=True)

    new_xml = f'{to_location}\\{abb_name}_{year}_RandomForest.tif.xml'

    #because this xml file is edited so heavily, we should be able to grab any xml as a template
    #right now, it seems like if we always take a single template xml we can prevent most issues 
    copy(template_xml,  new_xml)

    #rename all files to the LOCATION_YYYY_RandomForest convention when copying to the x staging folder
    for f in from_location.glob('*.*'):
        extension = '.'.join(f.name.split('.')[1:])
        if '.doc' not in extension:
            new_name = f'{abb_name}_{year}_RandomForest.{extension}'
            new_file = f'{to_location}/{new_name}'
            copy(f, new_file)

setupDirectories()

#----------------metadata elements------------------
def getReportingDatasets(root = root_path):
    '''Get all of the datasets used in classification as strings.'''

    reporting_folder = []
    folders = [f.name for f in root.glob('**/*') if f.is_dir()] 
    for n in folders:
        if 'eporting' in n:
            reporting_folder.append(n)

    target_folder = reporting_folder[-1]
    #the new method is to get reporting info in a json, but the old format is just a word doc
    #this accounts for both methods automatically
    jsons = [i for i in Path(root / target_folder).glob('*.json')]
    if jsons[0] != None:
        classification_stats = json.dumps(open(jsons[0]))
        used_datasets = classification_stats['datasets']
    else:
        reporting_doc = Document(root / target_folder / f'{area}-{year}-v{target_folder.split("V")[-1]}-classification_Irrigated_lands_reporting.docx')

        #get the datasets used in processing
        doc_metadata_table = reporting_doc.tables[0]
        rows = []
        for report in doc_metadata_table.column_cells(0):
            rows.append(report.text)
            if 'dataset' in report.text or 'Dataset' in report.text:
                column_index = rows.index(report.text)

        #the list of datasets we used in classification NOTE: this currently does not include datasets used to post process
        used_datasets = doc_metadata_table.cell(column_index, 1).text.strip("[]").replace("'", "").split(', ')

    #empty lists to be filled later 
    description_datasets = []
    reference_datasets = []
    post_process_datasets = []

    #make individual lists for the datasets we used and their references 
    for band in used_datasets:
        dataset = datasets_dict[band]
        ref_num = used_datasets.index(band) + 2 #there are two references that are always present, hence the +2

        if band in ['USDA/NASS/CDL', 'USDA/NAIP/DOQQ']:
            name = f'{dataset[0]}({ref_num})'
            post_process_datasets.append(name)
        else:
            name = f'{dataset[0]}({ref_num})'
            description_datasets.append(name)

        reference = f'({ref_num}) {dataset[1]}'
        reference_datasets.append(reference)

    datasets = ', '.join(description_datasets)
    references = '\n\n'.join(reference_datasets)
    post_process = ', '.join(post_process_datasets)

    return datasets, references, post_process

formatted_datasets, formatted_references, formatted_post_process = getReportingDatasets()

 
def updateMetadataDoc(metadata = metadata_doc, full = full_name, abb = abb_name, year = year,
                      datasets = formatted_datasets, references=formatted_references,
                      post_process= formatted_post_process):
    '''Get the metadata doc, replace keywords with new text, and create a dictionary of sections to pull later'''
    #a dictionary of how to update the document text to make sure it is matching the correct values
    dict = {'Region full': full,
            'Region abv.': abb,
            'Year': year,
            'Datasets': datasets,
            'References': references,
            'Post Process': post_process}
    docx_replace(doc=metadata, **dict)

    #the sections of the metadata document to parse, sections identified by text style in the word doc
    sections = {'Summary':[], 'Description':[], 'Normal':[], 
                'Credits':[], 'Use limitations':[], 'Extras':[]}

    #a loop to grab the strings from the updated document, get them into a single string, and fill out the dictionary
    for s in sections.keys():
        for p in metadata.paragraphs:
            # Filter for standard styles
            if p.style.name.startswith(s):
                sections[s].append(p.text)
        sections[s] = '\n\n'.join(sections[s])
    return sections

sections = updateMetadataDoc()

#the actual values the metadata will be updated with 
TAC = sections['Use limitations']
tags = f'Supervised Land Classification, Machine Learning, Random Forest, Water Budget, Monitoring, Hydrology, Groundwater, Surface Water, Irrigated Areas, Irrigation, Irrigated, Non-Irrigated, Regulatory, Farming, Idaho, ID, Idaho, Water, Water Use, {full_name}, {abb_name}, IDWR GIS Department, Environment'
place_keywords = f'{full_name}, {abb_name}, Idaho, ID'
file_title = f'{year} Irrigated Lands for the {full_name} ({abb_name}): Machine Learning Generated'
complete_file_name = f'{abb_name}_{year}_RandomForest'
description = sections['Description']
summary = sections['Summary']
extras = sections['Extras']
#to preserve the original thumbnail after arcpy steals it, I'm creating two links 
original_thumbnail_link = Path(temp_folder) / 'thumbnail.png'
thumbnail_link = Path(temp_folder) / 'thumbnail_backup.png' #NOTE: for whatever reason, this seems to cause problems. The file keeps moving after the code runs, which stops the code from running, but the code runs fine if you run it twice.
copy(original_thumbnail_link, thumbnail_link)         #which is why this line is here, to hedge our bets 
creation_date = strftime('%Y-%m-%d %H:%M:%S', strptime(ctime(getmtime(str(list(Path(n_loc).glob('*.tif'))[0])))))
publication_date = strftime('%Y-%m-%d %H:%M:%S', strptime(ctime(time())))
edition_date = strftime('%Y-%m-%d', strptime(ctime(time())))

def makeShortDate(date):
    '''change the date format to a short format (YYYYMMdd)'''
    return strftime('%Y%m%d', strptime(date, '%Y-%m-%d %H:%M:%S'))
creation_short = makeShortDate(creation_date)
publication_short = makeShortDate(publication_date)
#----------------defining metadata----------------------
target_tif =[i for i in Path(x_staging_loc).glob('*.tif')][0]
target_tif_meta = metadata.Metadata(target_tif)

#updating of metadata pieces through the ESRI interface
target_tif_meta.title = file_title
target_tif_meta.accessConstraints = TAC
target_tif_meta.tags = tags
target_tif_meta.credits = 'Idaho Department of Water Resources (IDWR)'
#this is a failsafe in case arcpy consumes the thumbnail. it seems like just rerunning solves the issue, so trying that here
try:
    target_tif_meta.thumbnailUri = str(thumbnail_link)
except:
    target_tif_meta.thumbnailUri = str(original_thumbnail_link)
target_tif_meta.description = description
target_tif_meta.summary = summary

#this coming section updates dates within the xml file because ESRI does not have a built in md method for dates. 
# It is, however, using the arc metadata template from the metadata object
xml_string = target_tif_meta.xml
root = ET.fromstring(xml_string)

#NOTE: we can probaly edit metadata in a few places, or omit, or some things need edits. Here is a list: (looking at ESPA 2024 for reference)
#DONE NEEDS REVIEW'Extents: Description' should be about the extent,
#TODO'Resource Contstraints' whole thing can be the same as access constraints or does it need to be different?,
#DONE, SHOULD GET REVIEWED 'Data Quality: Data Qaulity Report - Conceptual consistency' should proabably be something related to how we do QAQC for both measure reference and procedure
#DONE, SHOULD GET REVIEWED'Data Quality: Data Qaulity Report - Completeness Omission' I'm not sure what this is or hwo it is different than above so needs researched,
#TODO Pretty much everything in 'Data Quality' Needs more intel before we know what to put in there,
#TODO'Lineage: Lineage statment' Probably can be the general steps and tools we use to make the dataset maybe?,
#TODO'Lineage: Process Step' I dont know if this is just the last step or a specific line out of steps or how it relates to rationale,
#TODO'Geoprocessing History' I imagine this is the same as lineage, but it looks like it is just tracking the last tool used on the dataset in Arc,
#TODO'Fields': I dont know if we can or should alter any of this, except for changing how it references the .vat file, as it is referencing the wrong dataset
#TODO'Metadata Details': Is this similar to the lineage? I don't know why it is referencing such an old dataset, but also don't know if we can edit that
#TODO'Metadata Constraints':these are again constraints and limitations that may need to be refined for this section, but they may also be fine to mirror the other constraints

other_metadata ={'.//Esri/CreaDate': creation_short,
                 './/Esri/ModDate': publication_short, 
                 './/Esri/SyncDate': publication_short,
                 './/mdDateSt': creation_short,
                 './/dataIdInfo/idCitation/date': None,
                 './/dataIdInfo/idCitation/date/pubDate': publication_date,
                 './/dataIdInfo/idCitation/date/createDate': creation_date,
                 './/dataIdInfo/idCitation/resEdDate': edition_date,
                 './/dqInfo/dataLineage/prcStep/stepDateTm': publication_date,
                 './/eainfo/detailed/attr/begdatea': f'{year}0301',
                 './/eainfo/detailed/attr/enddatea': f'{year}1101',
                 './/dataIdInfo/placeKeys': place_keywords,
                 './/dataIdInfo/placeKeys/keyword': place_keywords,
                 './/dataIdInfo/tempKeys/keyword': year,
                 './/dataIdInfo/idCitation/resAltTitle': f'RF_IrrigatedLands_{year}_{abb_name}',
                 './/Esri/DataProperties/itemProps/itemName': f'{abb_name}_{year}_RandomForest.tif',
                 './/dataIdInfo/idCitation/datasetSeries/seriesName': file_title.split(' ', 1)[1],
                 './/Esri/DataProperties/itemProps/itemLocation': f'{x_spatial_loc}\\{abb_name}_{year}_RandomForest.tif',
                 './/Esri/DataProperties/itemProps/itemLocation/linkage':f'{x_spatial_loc}\\{abb_name}_{year}_RandomForest.tif',
                 './/dqInfo/dataLineage/prcStep/stepDesc': None, #NOTE: It seems like these are just steps used to create the dataset, do we need to provide this or do the methods kind of sum it up?
                 './/dqInfo/dataLineage/prcStep/stepRat': None,
                 './/dqInfo/dataLineage/statement': summary,
                 './/dataIdInfo/idCitation/otherCitDet': None,
                 './/dqInfo/report[@type="DQConcConsis"]/measDesc': long_metadata['DQC_measDesc'], 
                 './/dqInfo/report[@type="DQConcConsis"]/evalMethDesc': long_metadata['DQC_evalMethDesc'], 
                 './/dqInfo/report[@type="DQCompOm"]/measDesc': None,
                 './/dqInfo/report[@type="DQCompOm"]/evalMethDesc': None,
                 './/dqInfo/report[@type="DQQuanAttAcc"]/measDesc': long_metadata['DQQ_measDesc'],
                 './/dataIdInfo/idCitation/datasetSeries/issId': year,
                 './/dataIdInfo/dataExt/exDesc': str(long_metadata['exDesc']).format(full_name = full_name),
                 './/dataIdInfo/dataExt/tempEle/TempExtent/exTemp/TM_Period/tmBegin': f'{year}-03-01T00:00:00',
                 './/dataIdInfo/dataExt/tempEle/TempExtent/exTemp/TM_Period/tmEnd': f'{year}-11-01T00:00:00',
                 './/dataIdInfo/tpCat': 'Irrigated Lands',
                 './/dataIdInfo/tpCat/TopicCatCd': 'Irrigated Lands',
                 './/eainfo/detailed/enttyp/enttypl': f'{abb_name}_{year}_RandomForest.tif',
                 }

#this looks for every item in the above dictionary in the root xml string, changes it, then finally updates the target xml file
for d in other_metadata:
    el = root.find(d)
    el.text = other_metadata[d]
target_tif_meta.xml = ET.tostring(root, encoding='unicode')

if not target_tif_meta.isReadOnly:
    target_tif_meta.save()

print(f'New metadata saved to the tif at {str(target_tif)}')

#take the important info and store it in a json for the publishing script to access
publishing_json = {"file_title": file_title, 
                   "TAC": TAC, 
                   "tags": tags, 
                   "description": description, 
                   "temp_folder": temp_folder, 
                   "summary": summary, 
                   "thumbnail_link": thumbnail_link, 
                   "x_drive_name": x_drive_name, 
                   "year": year, 
                   "abb_name": abb_name, 
                   "x_staging_loc": x_staging_loc}
with open(Path(temp_folder)/'publishing_json.json', 'w') as f:
    json.dump(publishing_json, f)

#after you run this script, there are some items that GIS admin need to take care of, then when data is all set up on X:/Spatial 
# you run item_publishing.py
#%%
#helpers for checking your work without having to open arc catalog
use_helpers = False
if use_helpers:
    for elem in root.iter():
        if 'date' in elem.tag.lower() or 'Date' in elem.tag:
            # build the path from root down to this element
            path = []
            e = elem
            # ElementTree doesn't track parents, so just print tag + text for now
            print(elem.tag, '->', elem.text)

    #a helper loop that finds instances of a string in the xml and prints the path to update the dicitonary with if needed
    parent_map = {c: p for p in root.iter() for c in p}

    def get_path(elem):
        path = [elem.tag]
        while elem in parent_map:
            elem = parent_map[elem]
            path.append(elem.tag)
        return '/'.join(reversed(path))

    for elem in root.iter():
        if 'enttypl' in elem.tag.lower():
            print(get_path(elem), '=', elem.text)

#%%

#helper to view the xml in case you need to find a path for editing metadata 
if use_helpers:
    import xml.dom.minidom as minidom
    
    tgt_md = metadata.Metadata(target_tif)
    pretty = minidom.parseString(tgt_md.xml).toprettyxml(indent='  ')
    print(pretty)

